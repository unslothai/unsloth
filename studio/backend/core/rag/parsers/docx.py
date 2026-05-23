# SPDX-License-Identifier: AGPL-3.0-only
# Copyright 2026-present the Unsloth AI Inc. team. All rights reserved. See /studio/LICENSE.AGPL-3.0

from __future__ import annotations

from pathlib import Path

from . import ParsedPage


def extract(path: Path) -> list[ParsedPage]:
    from docx import Document

    document = Document(str(path))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text and paragraph.text.strip():
            parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n\n".join(parts).strip()
    if not text:
        return []
    return [ParsedPage(text = text, page_number = None)]
