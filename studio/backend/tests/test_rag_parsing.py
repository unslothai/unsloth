# SPDX-License-Identifier: AGPL-3.0-only
# Copyright 2026-present the Unsloth AI Inc. team. All rights reserved. See /studio/LICENSE.AGPL-3.0

"""PDF text extraction: layout-aware Markdown (pymupdf4llm) with plain-text fallback."""

from __future__ import annotations

import pytest

pytest.importorskip("pymupdf")


def _table_pdf(path):
    import pymupdf

    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_textbox(pymupdf.Rect(40, 40, 550, 70), "Quarterly Results", fontsize = 16)
    rows = [
        ("Quarter", "Revenue", "Growth"),
        ("Q1", "$1.2M", "12%"),
        ("Q2", "$1.5M", "25%"),
    ]
    y = 90
    for r in rows:
        page.insert_textbox(pymupdf.Rect(40, y, 250, y + 20), r[0], fontsize = 11)
        page.insert_textbox(pymupdf.Rect(250, y, 400, y + 20), r[1], fontsize = 11)
        page.insert_textbox(pymupdf.Rect(400, y, 540, y + 20), r[2], fontsize = 11)
        y += 24
    doc.save(str(path))
    doc.close()


def test_pdf_extracts_markdown_table(tmp_path, monkeypatch):
    # With Markdown on, the layout is emitted as Markdown markup (heading, and a pipe table
    # where the extractor detects one) that flat get_text never produces.
    pytest.importorskip("pymupdf4llm")
    from core.rag import config, parsers

    monkeypatch.setattr(config, "PDF_MARKDOWN", True)
    pdf = tmp_path / "table.pdf"
    _table_pdf(pdf)
    text = "\n".join(p.text for p in parsers.parse(str(pdf)))
    assert "Q2" in text and "$1.5M" in text  # cell values preserved
    assert "#" in text or "|" in text  # Markdown markup (heading or table pipes)


def test_pdf_markdown_off_uses_plain_text(tmp_path, monkeypatch):
    # The toggle (RAG_PDF_MARKDOWN=0) falls back to flat PyMuPDF text: content is still
    # there, but with no Markdown markup.
    from core.rag import config, parsers

    monkeypatch.setattr(config, "PDF_MARKDOWN", False)
    pdf = tmp_path / "table.pdf"
    _table_pdf(pdf)
    text = "\n".join(p.text for p in parsers.parse(str(pdf)))
    assert "Q2" in text and "$1.5M" in text
    assert (
        "#" not in text and "|" not in text
    )  # plain text path emits no Markdown markup


def test_pdf_bytes_use_same_extraction_path(tmp_path, monkeypatch):
    from core.rag import config, parsers

    monkeypatch.setattr(config, "PDF_MARKDOWN", False)
    pdf = tmp_path / "table.pdf"
    _table_pdf(pdf)
    from_file = parsers.parse(str(pdf))
    from_bytes, total_pages = parsers.parse_pdf_bytes(pdf.read_bytes())
    assert [page.text for page in from_bytes] == [page.text for page in from_file]
    assert total_pages == len(from_file)


def test_pdf_bytes_limit_pages_before_extraction(monkeypatch):
    import pymupdf

    from core.rag import config, parsers

    monkeypatch.setattr(config, "PDF_MARKDOWN", False)
    doc = pymupdf.open()
    for marker in ("page one", "page two", "page three"):
        page = doc.new_page()
        page.insert_text((40, 40), marker)
    data = doc.tobytes()
    doc.close()

    pages, total_pages = parsers.parse_pdf_bytes(data, max_pages = 2)
    assert len(pages) == 2
    assert "page two" in pages[-1].text
    assert total_pages == 3  # full count, not the 2 extracted


def test_pdf_markdown_receives_page_limit(monkeypatch):
    from core.rag import parsers

    captured = {}

    class _FakePymupdf4llm:
        @staticmethod
        def to_markdown(doc, **kwargs):
            captured.update(kwargs)
            return [{"text": "page"} for _ in kwargs["pages"]]

    class _Doc:
        page_count = 100

    monkeypatch.setitem(__import__("sys").modules, "pymupdf4llm", _FakePymupdf4llm)
    assert parsers._pdf_markdown(_Doc(), range(2)) == ["page", "page"]
    assert captured == {"page_chunks": True, "show_progress": False, "pages": [0, 1]}


def test_pdf_markdown_passes_only_supported_legacy_kwargs(monkeypatch):
    # The pinned PyMuPDF4LLM legacy path ignores unknown kwargs; do not pass the
    # newer layout-only OCR knobs or Markdown extraction silently loses policy control.
    from core.rag import parsers

    captured = {}

    class _FakePymupdf4llm:
        @staticmethod
        def to_markdown(doc, **kwargs):
            captured.update(kwargs)
            return [{"text": "plain markdown"}]

    class _Doc:
        page_count = 1

    monkeypatch.setitem(__import__("sys").modules, "pymupdf4llm", _FakePymupdf4llm)
    assert parsers._pdf_markdown(_Doc()) == ["plain markdown"]
    assert captured == {"page_chunks": True, "show_progress": False}


def test_pdf_markdown_falls_back_when_lib_missing(tmp_path, monkeypatch):
    # If pymupdf4llm extraction returns None (missing/failed), parsing still yields the
    # plain-text pages rather than raising.
    from core.rag import config, parsers

    monkeypatch.setattr(config, "PDF_MARKDOWN", True)
    monkeypatch.setattr(parsers, "_pdf_markdown", lambda doc: None)
    pdf = tmp_path / "table.pdf"
    _table_pdf(pdf)
    pages = parsers.parse(str(pdf))
    assert pages and "Quarter" in pages[0].text


def _long_text_pdf(path):
    import pymupdf

    doc = pymupdf.open()
    page = doc.new_page()
    body = "The quick brown fox jumps over the lazy dog. " * 12  # >200 letters
    page.insert_textbox(pymupdf.Rect(40, 40, 550, 750), body, fontsize = 11)
    doc.save(str(path))
    doc.close()


def test_pdf_markdown_corruption_falls_back_to_plain(tmp_path, monkeypatch):
    # pymupdf4llm can emit shaped RTL Presentation Forms for Arabic/Hebrew; the parser
    # detects that and uses PyMuPDF's logical-order text instead of the mangled Markdown.
    from core.rag import config, parsers

    monkeypatch.setattr(config, "PDF_MARKDOWN", True)
    shaped = "".join(chr(c) for c in range(0xFE8D, 0xFEA0)) * 20  # heavy shaped forms
    monkeypatch.setattr(parsers, "_pdf_markdown", lambda doc: [shaped] * doc.page_count)
    pdf = tmp_path / "table.pdf"
    _table_pdf(pdf)
    text = "\n".join(p.text for p in parsers.parse(str(pdf)))
    assert "Quarter" in text  # real logical-order text recovered
    assert not parsers._markdown_corrupted(text)  # shaped garbage not carried through


def test_pdf_markdown_incomplete_falls_back_to_plain(tmp_path, monkeypatch):
    # If pymupdf4llm silently drops most of a page, the parser prefers the fuller raw layer.
    from core.rag import config, parsers

    monkeypatch.setattr(config, "PDF_MARKDOWN", True)
    monkeypatch.setattr(parsers, "_pdf_markdown", lambda doc: ["x"] * doc.page_count)
    pdf = tmp_path / "long.pdf"
    _long_text_pdf(pdf)
    text = "\n".join(p.text for p in parsers.parse(str(pdf)))
    assert (
        "quick brown fox" in text
    )  # fuller raw layer used, not the near-empty Markdown


def _docx_with_table(path):
    import docx

    document = docx.Document()
    document.add_paragraph("Intro before table.")
    table = document.add_table(rows = 2, cols = 2)
    table.cell(0, 0).text = "NAME"
    table.cell(0, 1).text = "SCORE"
    table.cell(1, 0).text = "Alice"
    table.cell(1, 1).text = "97pts"
    document.add_paragraph("Outro after table.")
    document.save(str(path))


def test_docx_extracts_table_cells(tmp_path):
    # document.paragraphs alone drops tables; the parser walks body content in order so
    # table cells survive (pipe-joined, which the preview locator anchors on).
    pytest.importorskip("docx")
    from core.rag import parsers

    docx_path = tmp_path / "t.docx"
    _docx_with_table(docx_path)
    text = "\n".join(p.text for p in parsers.parse(str(docx_path)))
    assert all(v in text for v in ("NAME", "SCORE", "Alice", "97pts"))  # cells kept
    assert "Alice | 97pts" in text  # row cells joined
    assert text.index("Intro") < text.index("NAME") < text.index("Outro")  # order kept


def test_docx_table_keeps_columns_and_collapses_cell_newlines(tmp_path):
    # Empty cells are kept (so columns stay aligned across rows) and a cell's internal
    # newlines are collapsed to spaces (so a multi-paragraph cell can't break the row).
    pytest.importorskip("docx")
    import docx

    from core.rag import parsers

    document = docx.Document()
    table = document.add_table(rows = 2, cols = 3)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = ""  # empty middle cell
    table.cell(0, 2).text = "C"
    multiline = table.cell(1, 0)
    multiline.text = "line1"
    multiline.add_paragraph("line2")  # cell now holds an internal newline
    table.cell(1, 1).text = "mid"
    table.cell(1, 2).text = "end"
    path = tmp_path / "aligned.docx"
    document.save(str(path))

    text = "\n".join(p.text for p in parsers.parse(str(path)))
    assert "A |  | C" in text  # empty cell preserved -> columns line up
    assert "line1 line2 | mid | end" in text  # internal newline collapsed to a space


def test_docx_table_merged_cell_keeps_grid_alignment(tmp_path):
    # A horizontally merged cell repeats across the spanned columns: emit its text once
    # then a placeholder, so the row keeps as many fields as its siblings (columns stay
    # aligned) without duplicating the merged text.
    pytest.importorskip("docx")
    import docx

    from core.rag import parsers

    document = docx.Document()
    table = document.add_table(rows = 2, cols = 3)
    table.cell(0, 0).text = "WIDE"
    table.cell(0, 2).text = "END"
    table.cell(0, 0).merge(table.cell(0, 1))  # span the first two columns
    table.cell(1, 0).text = "a"
    table.cell(1, 1).text = "b"
    table.cell(1, 2).text = "c"
    path = tmp_path / "merged.docx"
    document.save(str(path))

    text = "\n".join(p.text for p in parsers.parse(str(path)))
    assert text.count("WIDE") == 1  # merged cell not duplicated across spanned columns
    assert (
        "WIDE |  | END" in text
    )  # placeholder keeps 3 fields, aligned with "a | b | c"
    assert "a | b | c" in text


def test_docx_table_pads_omitted_grid_columns(tmp_path):
    # A row that skips leading grid columns exposes the gap via grid_cols_before; pad it
    # with empty fields so the value stays under the right header instead of shifting left.
    pytest.importorskip("docx")
    import docx
    from docx.oxml.ns import qn

    from core.rag import parsers

    document = docx.Document()
    table = document.add_table(rows = 2, cols = 3)
    table.cell(0, 0).text = "H1"
    table.cell(0, 1).text = "H2"
    table.cell(0, 2).text = "H3"
    tr = table.rows[1]._tr  # drop the first cell and mark it skipped via <w:gridBefore>
    tr.remove(tr.tc_lst[0])
    trPr = tr.get_or_add_trPr()
    trPr.insert(0, trPr.makeelement(qn("w:gridBefore"), {qn("w:val"): "1"}))
    table.rows[1].cells[0].text = "X"  # sits in column 2
    path = tmp_path / "gap.docx"
    document.save(str(path))

    text = "\n".join(p.text for p in parsers.parse(str(path)))
    assert " | X | " in text  # leading gap padded so X lines up under H2, not H1


def test_docx_flattens_nested_table(tmp_path):
    # cell.text ignores tables nested inside a cell; walk cell.tables so nested rows are
    # not silently dropped from the indexed text.
    pytest.importorskip("docx")
    import docx

    from core.rag import parsers

    document = docx.Document()
    outer = document.add_table(rows = 1, cols = 1).cell(0, 0)
    outer.text = "outer"
    nested = outer.add_table(rows = 1, cols = 2)
    nested.cell(0, 0).text = "NESTED-A"
    nested.cell(0, 1).text = "NESTED-B"
    path = tmp_path / "nested.docx"
    document.save(str(path))

    text = "\n".join(p.text for p in parsers.parse(str(path)))
    assert "NESTED-A | NESTED-B" in text  # nested table flattened, not dropped


def test_docx_nested_table_keeps_in_cell_order(tmp_path):
    # A cell holding paragraph, nested table, paragraph must serialize in that order
    # (cell.text alone would emit both paragraphs before the nested rows).
    pytest.importorskip("docx")
    import docx

    from core.rag import parsers

    document = docx.Document()
    cell = document.add_table(rows = 1, cols = 1).cell(0, 0)
    cell.text = "before"
    nested = cell.add_table(rows = 1, cols = 2)
    nested.cell(0, 0).text = "NESTED-A"
    nested.cell(0, 1).text = "NESTED-B"
    cell.add_paragraph("after")
    path = tmp_path / "nested_order.docx"
    document.save(str(path))

    text = "\n".join(p.text for p in parsers.parse(str(path)))
    assert text.index("before") < text.index("NESTED-A") < text.index("after")


def test_docx_table_vertical_merge_emitted_once(tmp_path):
    # A vertically merged cell maps every continuation row back to the origin <w:tc>;
    # emit it once and leave placeholders below so a row-spanning label isn't repeated.
    pytest.importorskip("docx")
    import docx

    from core.rag import parsers

    document = docx.Document()
    table = document.add_table(rows = 3, cols = 2)
    table.cell(0, 0).merge(table.cell(1, 0)).merge(table.cell(2, 0)).text = "SECTION"
    table.cell(0, 1).text = "r0"
    table.cell(1, 1).text = "r1"
    table.cell(2, 1).text = "r2"
    path = tmp_path / "vmerge.docx"
    document.save(str(path))

    text = "\n".join(p.text for p in parsers.parse(str(path)))
    assert text.count("SECTION") == 1  # not repeated on each spanned row
    assert "SECTION | r0" in text and " | r1" in text and " | r2" in text
