# SPDX-License-Identifier: AGPL-3.0-only
# Copyright 2026-present the Unsloth AI Inc. team. All rights reserved. See /studio/LICENSE.AGPL-3.0

"""Source preview + quick-edit routes: /documents/{id}/content (GET and PUT).

The properties worth defending here are the destructive ones. An edit is written
to Unsloth's managed uploads root and nowhere else; the document it replaces is
retired only after the re-index succeeds; and the replacement lands in the same
scope, so retrieval keeps finding the source where it was.
"""

from __future__ import annotations

import os
import time

import pytest

pytest.importorskip("sqlite_vec")

PROJECT_ID = "proj-1"


@pytest.fixture(autouse = True)
def chat_project(rag_home):
    """Both routes gate on the owning project still existing, via
    ``_require_document_owner``, so the scope needs a real chat_projects row."""
    from storage.studio_db import upsert_chat_project

    now = int(time.time() * 1000)
    upsert_chat_project({"id": PROJECT_ID, "name": "Test", "createdAt": now, "updatedAt": now})


def _client():
    from fastapi import FastAPI
    from fastapi.testclient import TestClient

    from auth.authentication import get_current_subject
    from routes.rag import router

    app = FastAPI()
    app.include_router(router, prefix = "/api/rag")
    app.dependency_overrides[get_current_subject] = lambda: "tester"
    return TestClient(app)


def _await_job(job_id: str, *, expect: str = "completed") -> dict:
    from core.rag import ingestion

    deadline = time.time() + 30
    status = None
    while time.time() < deadline:
        status = ingestion.get_job_status(job_id)
        if status and status["status"] in ("completed", "failed"):
            break
        time.sleep(0.02)
    assert status and status["status"] == expect, status
    return status


def _ingest(
    filename: str,
    body: str,
    *,
    project_id: str = "proj-1",
):
    """Index one managed upload into a project scope, as the upload route does."""
    from core.rag import ingestion, store
    from utils.paths import ensure_dir, rag_uploads_root

    stored = ensure_dir(rag_uploads_root()) / filename
    stored.write_text(body, encoding = "utf-8")
    scope = store.project_scope(project_id)
    doc_id, job_id = ingestion.start_ingestion(
        scope, None, None, filename, str(stored), project_id = project_id
    )
    _await_job(job_id)
    return scope, doc_id, str(stored)


def _document_row(document_id: str):
    from core.rag import store
    from storage import rag_db

    conn = rag_db.get_connection()
    try:
        return store.get_document(conn, document_id)
    finally:
        conn.close()


def _search(
    client,
    query: str,
    project_id: str = "proj-1",
):
    return client.post(
        "/api/rag/search",
        json = {"query": query, "project_id": project_id, "mode": "lexical"},
    ).json()["results"]


# GET /content


def test_markdown_source_is_returned_editable(rag_home, stub_embeddings):
    _, doc_id, _ = _ingest("notes.md", "# Notes\n\nkickoff is on the third\n")
    body = _client().get(f"/api/rag/documents/{doc_id}/content").json()

    assert body["mediaKind"] == "text"
    assert body["preview"] == "markdown"
    assert "kickoff is on the third" in body["text"]
    assert body["editable"] is True
    assert body["readOnlyReason"] is None
    assert body["truncated"] is False


def test_plain_text_source_opens_straight_in_the_editor(rag_home, stub_embeddings):
    # No richer view than the text itself, so the modal offers no View/Edit pair.
    _, doc_id, _ = _ingest("log.txt", "plain body\n")
    body = _client().get(f"/api/rag/documents/{doc_id}/content").json()

    assert body["preview"] == "source"
    assert body["editable"] is True


def test_html_previews_in_the_canvas_and_stays_editable(rag_home, stub_embeddings):
    # "html" routes the View tab to the sandboxed artifact canvas; Edit keeps the
    # raw markup, so the pair behaves like markdown's.
    markup = "<!DOCTYPE html><html><body><h1>Report</h1></body></html>"
    _, doc_id, _ = _ingest("page.html", markup)
    body = _client().get(f"/api/rag/documents/{doc_id}/content").json()

    assert body["preview"] == "html"
    assert body["text"] == markup, "Edit must show the markup, not stripped text"
    assert body["editable"] is True


def test_docx_preview_mode_is_extracted(rag_home, stub_embeddings):
    pytest.importorskip("docx")
    import docx

    from core.rag import ingestion, store
    from utils.paths import ensure_dir, rag_uploads_root

    path = ensure_dir(rag_uploads_root()) / "brief.docx"
    document = docx.Document()
    document.add_paragraph("body text")
    document.save(str(path))
    doc_id, job_id = ingestion.start_ingestion(
        store.project_scope(PROJECT_ID),
        None,
        None,
        "brief.docx",
        str(path),
        project_id = PROJECT_ID,
    )
    _await_job(job_id)

    assert _client().get(f"/api/rag/documents/{doc_id}/content").json()["preview"] == "extracted"


def test_pdf_is_display_only_and_carries_no_text(rag_home, stub_embeddings):
    # A PDF has no faithful plain-text round trip, so the modal renders the file
    # itself through the signed URL rather than offering an editor.
    pytest.importorskip("pymupdf")
    import pymupdf

    from core.rag import ingestion, store
    from utils.paths import ensure_dir, rag_uploads_root

    path = ensure_dir(rag_uploads_root()) / "paper.pdf"
    pdf = pymupdf.open()
    pdf.new_page().insert_text((72, 72), "attention is all you need", fontsize = 11)
    pdf.save(str(path))
    pdf.close()
    doc_id, job_id = ingestion.start_ingestion(
        store.project_scope("proj-1"), None, None, "paper.pdf", str(path), project_id = "proj-1"
    )
    _await_job(job_id)

    body = _client().get(f"/api/rag/documents/{doc_id}/content").json()
    assert body["mediaKind"] == "pdf"
    assert body["text"] is None
    assert body["editable"] is False
    assert "PDF" in body["readOnlyReason"]


def test_docx_shows_extracted_text_and_refuses_editing(rag_home, stub_embeddings):
    # Read-only for the same reason as a PDF, but it still previews: what it shows
    # is exactly the text that was chunked and embedded.
    pytest.importorskip("docx")
    import docx

    from core.rag import ingestion, store
    from utils.paths import ensure_dir, rag_uploads_root

    path = ensure_dir(rag_uploads_root()) / "brief.docx"
    document = docx.Document()
    document.add_paragraph("quarterly revenue is up twelve percent")
    document.save(str(path))
    doc_id, job_id = ingestion.start_ingestion(
        store.project_scope("proj-1"), None, None, "brief.docx", str(path), project_id = "proj-1"
    )
    _await_job(job_id)

    body = _client().get(f"/api/rag/documents/{doc_id}/content").json()
    assert body["mediaKind"] == "text"
    assert "quarterly revenue is up twelve percent" in body["text"]
    assert body["editable"] is False
    assert (
        body["readOnlyReason"]
        == "Word documents are shown as the text indexed and cannot be edited here."
    )


def test_oversized_text_previews_truncated_and_unsavable(rag_home, stub_embeddings, monkeypatch):
    # Truncating for display is fine; saving a truncated body would delete the
    # tail of the user's file, so editing is refused instead.
    from routes import rag as rag_routes

    monkeypatch.setattr(rag_routes, "_MAX_TEXT_EDIT_BYTES", 64)
    _, doc_id, _ = _ingest("big.md", "x" * 500)

    body = _client().get(f"/api/rag/documents/{doc_id}/content").json()
    assert body["truncated"] is True
    assert body["editable"] is False
    assert len(body["text"]) == 64


def test_content_of_an_unknown_document_is_404(rag_home, stub_embeddings):
    assert _client().get("/api/rag/documents/nope/content").status_code == 404


def test_content_refuses_a_stored_path_outside_the_uploads_root(
    rag_home, stub_embeddings, tmp_path
):
    # Defense in depth, mirroring the signed-file route: a row whose stored_path
    # escaped the managed root is never read.
    from storage import rag_db

    _, doc_id, _ = _ingest("notes.md", "body\n")
    outside = tmp_path / "outside.md"
    outside.write_text("secret", encoding = "utf-8")
    conn = rag_db.get_connection()
    try:
        conn.execute("UPDATE documents SET stored_path=? WHERE id=?", (str(outside), doc_id))
        conn.commit()
    finally:
        conn.close()

    assert _client().get(f"/api/rag/documents/{doc_id}/content").status_code == 403


# PUT /content


def test_saving_an_edit_reindexes_into_the_same_scope(rag_home, stub_embeddings):
    """The retrieval-safety case: after an edit the model finds the new text, in
    the same project, and the superseded document is gone."""
    scope, doc_id, old_path = _ingest("notes.md", "kickoff is on the third\n")
    client = _client()
    assert _search(client, "kickoff third")

    res = client.put(
        f"/api/rag/documents/{doc_id}/content",
        json = {"text": "kickoff moved to the seventeenth\n"},
    )
    assert res.status_code == 200, res.text
    new_id = res.json()["documentId"]
    assert new_id != doc_id, "an edit is ingested as a replacement document"
    _await_job(res.json()["jobId"])

    row = _document_row(new_id)
    assert row["scope"] == scope
    assert row["project_id"] == "proj-1"
    assert row["filename"] == "notes.md"

    assert _search(client, "seventeenth"), "the edited text is not retrievable"
    assert _document_row(doc_id) is None, "the superseded document was not retired"
    assert not _search(client, "third"), "the pre-edit text is still indexed"


def test_saving_writes_only_inside_the_uploads_root(rag_home, stub_embeddings):
    from utils.paths import rag_uploads_root

    _, doc_id, old_path = _ingest("notes.md", "before\n")
    client = _client()
    res = client.put(f"/api/rag/documents/{doc_id}/content", json = {"text": "after\n"})
    _await_job(res.json()["jobId"])

    new_path = _document_row(res.json()["documentId"])["stored_path"]
    uploads = os.path.realpath(str(rag_uploads_root()))
    assert os.path.commonpath([uploads, os.path.realpath(new_path)]) == uploads
    assert new_path != old_path, "the edit must not overwrite the file being replaced"
    assert not os.path.exists(old_path), "the replaced upload was left behind"
    with open(new_path, encoding = "utf-8") as handle:
        assert handle.read() == "after\n"


def test_a_failed_reindex_leaves_the_original_searchable(rag_home, stub_embeddings, monkeypatch):
    """`replaces` retires the old document only on success. If the re-index dies,
    the user still has the source they started with."""
    from core.rag import embeddings

    _, doc_id, old_path = _ingest("notes.md", "kickoff is on the third\n")
    client = _client()

    working = embeddings.encode

    def boom(*args, **kwargs):
        raise RuntimeError("embedder exploded")

    monkeypatch.setattr(embeddings, "encode", boom)
    res = client.put(f"/api/rag/documents/{doc_id}/content", json = {"text": "replacement\n"})
    assert res.status_code == 200
    _await_job(res.json()["jobId"], expect = "failed")

    assert _document_row(doc_id) is not None, "the original was destroyed by a failed edit"
    assert os.path.exists(old_path), "the original's file was removed by a failed edit"
    # Restore only this patch: monkeypatch.undo() would also drop the stubbed
    # embedder the fixture installed, sending the search below to a real model.
    monkeypatch.setattr(embeddings, "encode", working)
    assert _search(client, "kickoff third"), "the original is no longer retrievable"


def test_editing_html_saves_the_markup_and_indexes_its_visible_text(rag_home, stub_embeddings):
    # The Edit tab holds markup, so the file must round-trip as markup while the
    # index keeps holding the stripped text the HTML parser produces.
    _, doc_id, _ = _ingest("page.html", "<html><body><p>before</p></body></html>")
    client = _client()

    edited = "<html><body><p>afterwards indexed</p><script>ignored()</script></body></html>"
    res = client.put(f"/api/rag/documents/{doc_id}/content", json = {"text": edited})
    assert res.status_code == 200
    new_id = res.json()["documentId"]
    _await_job(res.json()["jobId"])

    reopened = client.get(f"/api/rag/documents/{new_id}/content").json()
    assert reopened["text"] == edited, "the markup did not round-trip"
    assert _search(client, "afterwards indexed")
    assert not _search(client, "ignored"), "script contents must not be indexed"


def test_emptying_a_source_keeps_it_as_an_empty_document(rag_home, stub_embeddings):
    # Clearing the textarea is reachable by accident, so it must land somewhere
    # sane: the source survives with no chunks rather than the save half-failing
    # and leaving both versions behind.
    _, doc_id, _ = _ingest("notes.md", "kickoff is on the third\n")
    client = _client()

    res = client.put(f"/api/rag/documents/{doc_id}/content", json = {"text": "   \n"})
    assert res.status_code == 200
    new_id = res.json()["documentId"]
    _await_job(res.json()["jobId"])

    row = _document_row(new_id)
    assert row["status"] == "completed"
    assert row["num_chunks"] == 0
    assert _document_row(doc_id) is None, "the superseded document was not retired"
    assert not _search(client, "kickoff third")


def test_a_linked_folder_source_cannot_be_edited(rag_home, stub_embeddings, tmp_path):
    """The hard constraint: nothing in the user's own folder is ever written to,
    and its snapshot is not editable either (a sync would wipe the edit)."""
    from storage import rag_db

    original = tmp_path / "linked.md"
    original.write_text("owned by the folder\n", encoding = "utf-8")
    before = original.read_bytes()
    before_mtime = original.stat().st_mtime_ns

    _, doc_id, _ = _ingest("linked.md", "owned by the folder\n")
    conn = rag_db.get_connection()
    try:
        conn.execute("UPDATE documents SET linked_folder_id=? WHERE id=?", ("folder-1", doc_id))
        conn.commit()
    finally:
        conn.close()

    client = _client()
    assert client.get(f"/api/rag/documents/{doc_id}/content").json()["editable"] is False
    res = client.put(f"/api/rag/documents/{doc_id}/content", json = {"text": "rewritten\n"})
    assert res.status_code == 409

    assert original.read_bytes() == before, "the linked original was modified"
    assert original.stat().st_mtime_ns == before_mtime, "the linked original was touched"
    assert _document_row(doc_id) is not None


def test_a_pdf_cannot_be_saved(rag_home, stub_embeddings):
    pytest.importorskip("pymupdf")
    import pymupdf

    from core.rag import ingestion, store
    from utils.paths import ensure_dir, rag_uploads_root

    path = ensure_dir(rag_uploads_root()) / "paper.pdf"
    pdf = pymupdf.open()
    pdf.new_page().insert_text((72, 72), "body", fontsize = 11)
    pdf.save(str(path))
    pdf.close()
    doc_id, job_id = ingestion.start_ingestion(
        store.project_scope("proj-1"), None, None, "paper.pdf", str(path), project_id = "proj-1"
    )
    _await_job(job_id)

    res = _client().put(f"/api/rag/documents/{doc_id}/content", json = {"text": "nope"})
    assert res.status_code == 400
    assert os.path.exists(path), "a rejected save must not disturb the stored file"


def test_a_source_that_is_still_indexing_cannot_be_saved(rag_home, stub_embeddings):
    # A live ingestion worker owns the file and is writing this document's rows.
    from storage import rag_db

    _, doc_id, _ = _ingest("notes.md", "body\n")
    conn = rag_db.get_connection()
    try:
        conn.execute("UPDATE documents SET status='running' WHERE id=?", (doc_id,))
        conn.commit()
    finally:
        conn.close()

    res = _client().put(f"/api/rag/documents/{doc_id}/content", json = {"text": "edited"})
    assert res.status_code == 409


def test_saving_over_the_size_cap_is_rejected(rag_home, stub_embeddings):
    from routes.rag import _MAX_TEXT_EDIT_BYTES

    _, doc_id, _ = _ingest("notes.md", "body\n")
    res = _client().put(
        f"/api/rag/documents/{doc_id}/content",
        json = {"text": "x" * (_MAX_TEXT_EDIT_BYTES + 1)},
    )
    assert res.status_code == 422


def test_a_file_with_invalid_utf8_is_shown_but_not_editable(rag_home, stub_embeddings):
    # Decoding with errors="replace" turns a stray byte into U+FFFD. Saving that
    # back would rewrite the byte as the replacement character, corrupting content
    # the user never touched, so the source is read-only.
    from core.rag import ingestion, store
    from utils.paths import ensure_dir, rag_uploads_root

    path = ensure_dir(rag_uploads_root()) / "latin.txt"
    path.write_bytes(b"caf\xe9 latte was served\n")
    doc_id, job_id = ingestion.start_ingestion(
        store.project_scope(PROJECT_ID),
        None,
        None,
        "latin.txt",
        str(path),
        project_id = PROJECT_ID,
    )
    _await_job(job_id)

    body = _client().get(f"/api/rag/documents/{doc_id}/content").json()
    assert "latte was served" in body["text"], "the file should still preview"
    assert body["editable"] is False
    assert body["truncated"] is True
    assert path.read_bytes().startswith(b"caf\xe9"), "the file itself is untouched"


def test_valid_utf8_multibyte_text_stays_editable(rag_home, stub_embeddings):
    # The guard above must not catch ordinary non-ASCII: a correctly encoded file
    # round-trips exactly and stays editable.
    _, doc_id, _ = _ingest("accents.md", "café, naïve, 日本語\n")
    body = _client().get(f"/api/rag/documents/{doc_id}/content").json()

    assert body["editable"] is True
    assert body["truncated"] is False
    assert "日本語" in body["text"]


def test_saving_is_capped_on_encoded_bytes_not_characters(rag_home, stub_embeddings):
    # Pydantic's max_length counts characters. A CJK payload under that limit can
    # still encode to more than the byte cap, and saving it would produce a file
    # the next GET truncates -- turning the source read-only right after saving.
    from routes.rag import _MAX_TEXT_EDIT_BYTES

    _, doc_id, _ = _ingest("notes.md", "body\n")
    oversized = "你" * (_MAX_TEXT_EDIT_BYTES // 3 + 1)  # 3 bytes each in UTF-8
    assert len(oversized) <= _MAX_TEXT_EDIT_BYTES, "must pass the character check"
    assert len(oversized.encode("utf-8")) > _MAX_TEXT_EDIT_BYTES

    res = _client().put(f"/api/rag/documents/{doc_id}/content", json = {"text": oversized})
    assert res.status_code == 413
    assert _document_row(doc_id) is not None, "the original must survive a refused save"


def test_a_second_concurrent_save_is_refused(rag_home, stub_embeddings):
    """Two saves of one source would each retire the same old row, leaving both
    replacements indexed. The first claims the document; the second loses."""
    _, doc_id, _ = _ingest("notes.md", "original\n")
    client = _client()

    first = client.put(f"/api/rag/documents/{doc_id}/content", json = {"text": "one\n"})
    assert first.status_code == 200
    # The claim is held until the replacement retires the row, so a save landing
    # while the first is still in flight is refused rather than racing it.
    second = client.put(f"/api/rag/documents/{doc_id}/content", json = {"text": "two\n"})
    assert second.status_code == 409

    _await_job(first.json()["jobId"])
    assert _document_row(doc_id) is None
    assert _search(client, "one")
    assert not _search(client, "two"), "the refused save must not have been indexed"


def test_a_refused_start_releases_the_claim(rag_home, stub_embeddings, monkeypatch):
    # If the replacement never starts, the source must not be left stuck reading
    # as indexing -- that would make it permanently uneditable.
    from core.rag import ingestion

    _, doc_id, _ = _ingest("notes.md", "body\n")
    client = _client()

    def refuse(*args, **kwargs):
        raise RuntimeError("no worker available")

    monkeypatch.setattr(ingestion, "start_ingestion", refuse)
    with pytest.raises(RuntimeError):
        client.put(f"/api/rag/documents/{doc_id}/content", json = {"text": "edited\n"})
        # Read the row before leaving the raising context: TestClient re-raises the
        # worker error on exit, and asserting afterwards runs against a torn-down
        # app rather than the state the release produced.
    row = _document_row(doc_id)
    assert (
        row is not None and row["status"] == "completed"
    ), "a replacement that never started must leave the source editable again"


def test_replaces_requires_dedupe_off(rag_home, stub_embeddings):
    """The guard on the one new parameter: the dedupe branch owns `replaces`, and
    its early return would silently drop a caller-supplied value."""
    from core.rag import ingestion, store
    from utils.paths import ensure_dir, rag_uploads_root

    path = ensure_dir(rag_uploads_root()) / "x.md"
    path.write_text("body\n", encoding = "utf-8")
    with pytest.raises(ValueError, match = "dedupe"):
        ingestion.start_ingestion(
            store.project_scope("proj-1"),
            None,
            None,
            "x.md",
            str(path),
            replaces = ("some-id", None),
        )
