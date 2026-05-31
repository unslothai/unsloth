"""Document parser tests — each format skipped if its lib is unavailable.

Phase 3A: parsers now return ParseResult (iterable over .pages) and
emit Markdown so the chunker can split on heading boundaries.
"""

import sys
from pathlib import Path

import pytest

REPO_ROOT = Path(__file__).resolve().parents[2]
STUDIO_BACKEND = REPO_ROOT / "studio" / "backend"
if str(STUDIO_BACKEND) not in sys.path:
    sys.path.insert(0, str(STUDIO_BACKEND))


def test_text_parser_utf8(tmp_path):
    from core.rag.parsers import parse

    file = tmp_path / "sample.txt"
    file.write_text("hello world\n\nsecond paragraph", encoding = "utf-8")
    result = parse(file)
    assert len(result) == 1
    assert "hello world" in result.pages[0].text
    assert "second paragraph" in result.pages[0].text
    assert result.images == []


def test_markdown_parser_preserves_headings(tmp_path):
    from core.rag.parsers import parse

    file = tmp_path / "sample.md"
    file.write_text("# Title\n\nBody text with **emphasis**.", encoding = "utf-8")
    result = parse(file)
    assert result.pages
    # Markdown passes through; heading marker preserved.
    assert "# Title" in result.pages[0].text


def test_unsupported_format_raises(tmp_path):
    from core.rag.parsers import UnsupportedFormatError, parse

    file = tmp_path / "weird.xyz"
    file.write_text("nope")
    with pytest.raises(UnsupportedFormatError):
        parse(file)


def test_html_parser_emits_markdown_headings(tmp_path):
    pytest.importorskip("bs4")
    pytest.importorskip("lxml")
    pytest.importorskip("markdownify")
    from core.rag.parsers import parse

    file = tmp_path / "sample.html"
    file.write_text(
        "<html><body>"
        "<script>alert(1)</script>"
        "<h1>Main Title</h1>"
        "<h2>Sub Section</h2>"
        "<p>visible text</p>"
        "<ul><li>one</li><li>two</li></ul>"
        "</body></html>",
        encoding = "utf-8",
    )
    result = parse(file)
    assert result.pages
    md = result.pages[0].text
    # markdownify: <h1> → '# ', <h2> → '## '
    assert "# Main Title" in md
    assert "## Sub Section" in md
    assert "visible text" in md
    # script content scrubbed
    assert "alert" not in md
    # list items become Markdown bullets
    assert "one" in md and "two" in md


def test_pdf_parser_extracts_pages(tmp_path):
    pytest.importorskip("pymupdf")
    pytest.importorskip("pymupdf4llm")
    from core.rag.parsers import parse

    pypdf = pytest.importorskip("pypdf")
    from pypdf import PdfWriter

    file = tmp_path / "tiny.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width = 72, height = 72)
    with open(file, "wb") as f:
        writer.write(f)

    # Blank page: no text, returns empty pages without error.
    result = parse(file)
    assert isinstance(result.pages, list)
    assert isinstance(result.images, list)


def test_docx_parser_emits_markdown_headings(tmp_path):
    pytest.importorskip("docx")
    pytest.importorskip("mammoth")
    pytest.importorskip("markdownify")
    from docx import Document

    file = tmp_path / "sample.docx"
    doc = Document()
    doc.add_heading("Top Level Heading", level = 1)
    doc.add_paragraph("First paragraph here.")
    doc.add_heading("Sub Heading", level = 2)
    doc.add_paragraph("Second paragraph here.")
    doc.save(str(file))
    from core.rag.parsers import parse

    result = parse(file)
    assert result.pages
    md = result.pages[0].text
    # mammoth _STYLE_MAP: Heading 1/2 → h1/h2 → '# '/'## '.
    assert "# Top Level Heading" in md
    assert "## Sub Heading" in md
    assert "First paragraph" in md
    assert "Second paragraph" in md


def test_parse_result_is_iterable_for_backcompat(tmp_path):
    """Code that does `for page in parse(path)` should keep working."""
    from core.rag.parsers import parse

    file = tmp_path / "sample.txt"
    file.write_text("hello", encoding = "utf-8")
    result = parse(file)
    pages = list(result)
    assert len(pages) == 1
    assert pages[0].text == "hello"
